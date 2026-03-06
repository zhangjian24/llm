"""
Word 文档解析器
使用 python-docx 进行解析
"""
from docx import Document
from typing import Dict
from io import BytesIO
from .base_parser import DocumentParser
from app.exceptions import DocumentParseError


class DocxParser(DocumentParser):
    """
    Word 文档解析器
    
    特性:
    - 提取段落文本
    - 提取表格内容
    - 保留基本结构
    """
    
    @classmethod
    def get_supported_mime_types(cls) -> list[str]:
        return [
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ]
    
    async def parse(self, file_content: bytes) -> str:
        """
        解析 Word 文件为文本
        
        Args:
            file_content: Word 文件二进制内容
            
        Returns:
            str: 提取的文本内容
            
        Raises:
            DocumentParseError: 解析失败时抛出
        """
        try:
            # 读取文件
            file_io = BytesIO(file_content)
            doc = Document(file_io)
            
            # 提取段落文本
            text_parts = []
            
            # 提取所有段落
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text.strip())
            
            # 提取表格内容
            for table in doc.tables:
                table_text = []
                for row_idx, row in enumerate(table.rows):
                    row_cells = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_cells.append(cell.text.strip())
                    if row_cells:
                        table_text.append(" | ".join(row_cells))
                
                if table_text:
                    text_parts.append("\n".join(table_text))
            
            # 合并所有文本
            full_text = "\n\n".join(text_parts)
            
            return full_text.strip()
            
        except Exception as e:
            raise DocumentParseError(f"Word 解析失败：{str(e)}")
